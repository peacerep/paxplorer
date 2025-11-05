# server/actors.py
import io
import re
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from shiny import reactive, render, ui
from shinywidgets import output_widget, render_widget
import plotly.express as px
import plotly.graph_objects as go

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.png_export_utils import export_with_branding, get_data_version, LOGO_POSITIONS

# ---------------------------------------------------------------------
# CONSTANTS & DATA
# ---------------------------------------------------------------------
# DATA_DIR = Path(__file__).resolve().parent.parent / "data"
# # Data
# pax = pd.read_csv("data/pax.csv")
# con_ent = pd.read_csv("data/country_entity_geo_v9.csv")
# signatories = pd.read_csv("data/paax_signatory_v0.2_internal.csv").rename(columns={'stage_process':'stage_label'})
# pax_topics = pd.read_csv("data/all_pax_topics_no_imp.csv")
# pax_id_to_con = pd.read_csv("data/pax_id_to_con_info.csv")

from utils.data_loader import load_pax_data

data = load_pax_data()
pax = data["pax"]
pax_topics = data["pax_topics"]
pax_id_to_con = data["pax_id_to_con"]
con_ent = data["con_ent"]
signatories = data["signatories"].rename(columns={'stage_process': 'stage_label'})


# Stage & type color maps
stage_order = [
    "Pre-negotiation/process", "Ceasefire",
    "Framework-substantive, partial", "Framework-substantive, comprehensive",
    "Implementation", "Renewal", "Other"
]
stage_color_map = {
    "Pre-negotiation/process": "#016099",
    "Ceasefire": "#df1f36",
    "Framework-substantive, partial": "#fd8189",
    "Framework-substantive, comprehensive": "#fdd900",
    "Implementation": "#3aae2a",
    "Renewal": "#7b8ad6",
    "Other": "#c0de88",
}
type_color_map = {
    "Intrastate": "#0b5740",
    "Local": "#ac4399",
    "Interstate/mixed": "#df1f36",
    "Interstate": "#f28b20",
}

# Branding positions
ACTORS_LOGO_POSITION = (0.95, 0.97, 0.078, 0.078)
ACTORS_FILTER_TEXT_POSITION = (0.5, 0.004)
ACTORS_VERSION_POSITION = (0.98, 0.005)

SIGNATORY_UNIVERSE_TOTAL = signatories["AgtId"].nunique()

signatory_year_totals = (
    signatories.drop_duplicates(subset=["AgtId", "year"])
    .groupby("year")["AgtId"]
    .nunique()
    .reset_index(name="total")
)

# ---------------------------------------------------------------------
# SERVER
# ---------------------------------------------------------------------
def server(input, output, session):
    # -------------------------------------------------------
    #  HELPERS
    # -------------------------------------------------------
    def actor_title_prefix():
        # Prefer explicit actor names if any are selected
        names = input.actors_selected() or []
        if isinstance(names, str):
            names = [names]
        if names:
            shown = names[:2]
            if len(names) > 2:
                shown.append(f"+ {len(names)-2} more")
            return " / ".join(shown)

        # Else show the first selected type combo
        combos = input.actors_type_combo() or []
        if isinstance(combos, str):
            combos = [combos]
        if combos:
            return " | ".join(combos[:2]) + (f" + {len(combos)-2} more" if len(combos) > 2 else "")
        
        # Actor-attribute flags (international/regional/women)
        flags = input.actors_flags() or []
        flag_labels = {
            "international": "International",
            "regional": "Regional",
            "women": "Women",
        }
        if flags:
            # If only one attribute selected and no actor/type chosen, use that as title
            if not names and not combos and len(flags) == 1:
                return f"{flag_labels[flags[0]]} Actors"
            elif not names and not combos and len(flags) > 1:
                return " | ".join(flag_labels[f] for f in flags if f in flag_labels) + " Actors"

        return "All actors"

    def get_actor_filter_text():
        """One-line filter summary for PNG exports."""
        bits = []

        # --- Explicitly selected actor names ---
        names = input.actors_selected() or []
        if isinstance(names, str):
            names = [names]
        if names:
            show = names[:3]
            if len(names) > 3:
                show.append(f"+ {len(names)-3} more")
            bits.append("Actors: " + " | ".join(show))

        # --- Type/subtype combos ---
        combos = input.actors_type_combo() or []
        if isinstance(combos, str):
            combos = [combos]
        if combos:
            show = combos[:3]
            if len(combos) > 3:
                show.append(f"+ {len(combos)-3} more")
            bits.append("Actor type: " + " | ".join(show))

                # --- Actor Type selections ---
        selected_combos = input.actors_type_combo() or []
        if selected_combos:
            matching_actors_type = []
            for combo in selected_combos:
                combo = combo.strip()
                if ">" in combo:
                    main_type, rest = [p.strip() for p in combo.split(">", 1)]
                    subs = [p.strip() for p in rest.split("/") if p.strip()]
                else:
                    main_type = combo.strip()
                    subs = []

                cond = signatories["actor_type"] == main_type
                if subs:
                    cond &= signatories["sub_type"].isin(subs) | signatories["cs_type"].isin(subs)

                matching_actors_type.extend(signatories.loc[cond, "actor_name"].dropna().unique().tolist())

            matching_actors_type = sorted(set(matching_actors_type))
            if matching_actors_type:
                count_text = f" ({len(matching_actors_type)} actor{'s' if len(matching_actors_type) != 1 else ''})"
                bits.append(f"Actor Type: {', '.join(selected_combos)}{count_text}")


        # --- UN Type selections ---
        selected_un_types = input.actors_un_type() or []
        if selected_un_types:
            matching_actors = (
                signatories[signatories["un_type"].isin(selected_un_types)]["actor_name"]
                .dropna()
                .unique()
                .tolist()
            )
            if matching_actors:
                count_text = f" ({len(matching_actors)} actor{'s' if len(matching_actors) != 1 else ''})"
                bits.append(f"Actors: {', '.join(selected_un_types)}{count_text}")

        # --- Third-party only ---
        if input.actors_third_party_only():
            bits.append("Only showing third-party signatories")

        # --- General agreement filters ---
        reg = input.actors_region() or []
        cty = input.actors_country() or []
        pp = input.actors_peace_process() or []
        st = input.actors_stage() or []
        yrs = input.actors_year_range()

        if reg: bits.append("Region: " + ", ".join(reg))
        if cty: bits.append("Country: " + ", ".join(cty))
        if pp: bits.append("Processes: " + ", ".join(pp[:3]) + ("..." if len(pp) > 3 else ""))
        if st: bits.append("Stages: " + ", ".join(st))
        if yrs and len(yrs) == 2: bits.append(f"Years: {yrs[0]}–{yrs[1]}")

        return " | ".join(bits) if bits else "No filters applied"


    # -------------------------------------------------------
    #  ACTOR SELECTION LOGIC (MULTI + HIERARCHY)
    # -------------------------------------------------------
    @reactive.calc
    def actor_type_hierarchy_choices():
        """
        Build a dict of actor types -> list of readable sub/CS labels.
        All types appear even if there are no sub/CS combos.
        """
        df = signatories.copy()
        if df.empty:
            return {}

        # standardize blanks
        for col in ["actor_type", "sub_type", "cs_type"]:
            if col not in df.columns:
                df[col] = None

        choices = {}
        for t, group in df.groupby("actor_type", dropna=True):
            labels = set()
            # gather distinct (sub_type, cs_type) combos
            for _, row in group[["sub_type", "cs_type"]].drop_duplicates().iterrows():
                st = (str(row["sub_type"]).strip() if pd.notna(row["sub_type"]) else "")
                cs = (str(row["cs_type"]).strip() if pd.notna(row["cs_type"]) else "")
                parts = [p for p in [st, cs] if p]
                if parts:
                    labels.add(" / ".join(parts))
            choices[str(t).strip()] = sorted(labels)  # possibly empty
        return choices

    @reactive.effect
    def update_actor_type_combo():
        """Populate the actor type+sub/CS combo select (multi)."""
        flat = []
        for t, subs in actor_type_hierarchy_choices().items():
            if subs:
                for s in subs:
                    flat.append(f"{t} > {s}")
            else:
                flat.append(t)  # no dangling ">"
        ui.update_selectize(
            "actors_type_combo",
            choices=sorted(flat),
            selected=None,
            session=session,
        )

    @reactive.calc
    def filtered_actor_rows_for_name():
        """
        Return signatory rows filtered by selected type-combos and flags.
        Used to populate the actor_name choices.
        """
        df = signatories.copy()

        # Apply flags first (should work in isolation too)
        flags = input.actors_flags() or []
        for col in ["international", "regional", "women", "practical_third"]:
            if col in df.columns and col in flags:
                # accept truthy (1/"1"/True)
                df = df[df[col].fillna(0).astype(int) == 1]

        combos = input.actors_type_combo() or []
        if isinstance(combos, str):
            combos = [combos]

        if combos:
            mask = pd.Series(False, index=df.index)
            for combo in combos:
                # ✅ handle hierarchy: "Main > Sub / CS"
                combo = combo.strip()
                if ">" in combo:
                    main_type, rest = [p.strip() for p in combo.split(">", 1)]
                    subs = [p.strip() for p in rest.split("/") if p.strip()]
                else:
                    main_type = combo.strip()
                    subs = []

                cond = df["actor_type"] == main_type
                if subs:
                    cond &= df["sub_type"].isin(subs) | df["cs_type"].isin(subs)
                mask |= cond
            df = df[mask]

        return df


    @reactive.effect
    def update_actor_name_choices():
        """Populate multi-select actor_name choices; nothing selected by default."""
        df = filtered_actor_rows_for_name()
        names = sorted(df["actor_name"].dropna().unique()) if "actor_name" in df.columns else []
        ui.update_selectize(
            "actors_selected",
            choices=names,
            selected=[],
            session=session,
        )

    # -------------------------------------------------------
    #  GENERAL FILTERS (Region, Year, etc.)
    # -------------------------------------------------------
    @reactive.calc
    def region_choices(): return sorted(pax["Reg"].dropna().unique().tolist())

    @reactive.calc
    def country_choices(): return sorted(pax_id_to_con["name"].dropna().unique().tolist())

    @reactive.calc
    def agt_type_choices(): return sorted(pax["agt_type"].dropna().unique().tolist())

    @reactive.calc
    def peace_process_choices(): return sorted(pax["PPName"].dropna().unique().tolist())

    @reactive.calc
    def stage_choices(): return sorted(pax["stage_label"].dropna().unique().tolist())

    @reactive.calc
    def year_range(): return [int(pax["year"].min()), int(pax["year"].max())]

    @reactive.effect
    def populate_general_filters():
        ui.update_selectize("actors_region", choices=region_choices(), selected=[])
        ui.update_selectize("actors_country", choices=country_choices(), selected=[])
        ui.update_selectize("actors_agt_type", choices=agt_type_choices(), selected=[])
        ui.update_selectize("actors_peace_process", choices=peace_process_choices(), selected=[])
        ui.update_selectize("actors_stage", choices=stage_choices(), selected=[])
        yr = year_range()
        ui.update_slider("actors_year_range", min=yr[0], max=yr[1], value=yr)

    # -------------------------------------------------------
    #  DATA FILTERING PIPELINE
    # -------------------------------------------------------
    @reactive.calc
    def filtered_data_general():
        """Apply non-actor filters to pax data."""
        df = signatories.copy()

        countries = input.actors_country()
        if countries:
            agt_ids = pax_id_to_con[pax_id_to_con["name"].isin(countries)]["AgtId"].unique()
            df = df[df["AgtId"].isin(agt_ids)]

        agt_types = input.actors_agt_type()
        if agt_types:
            df = df[df["agt_type"].isin(agt_types)]

        yr = input.actors_year_range()
        if yr and len(yr) == 2:
            df = df[(df["year"] >= yr[0]) & (df["year"] <= yr[1])]

        regions = input.actors_region()
        if regions:
            df = df[df["Reg"].isin(regions)]

        pp = input.actors_peace_process()
        if pp:
            df = df[df["PPName"].isin(pp)]

        st = input.actors_stage()
        if st:
            df = df[df["stage_label"].isin(st)]

        return df

    @reactive.calc
    def filtered_agreements():
        """
        Combine general filters with actor-based filters (names, type-combos, flags).
        Flags apply even when nothing else is selected.
        """
        df = filtered_data_general()

        # Start with all signatories; apply flags first
        sig = signatories.copy()
        flags = input.actors_flags() or []
        for col in ["international", "regional", "women"]:
            if col in sig.columns and col in flags:
                sig[col] = sig[col].fillna(0).astype(int)
                sig = sig[sig[col] == 1]
        # --- New: restrict to third-party signatory instances if selected ---
        if input.actors_third_party_only():
            sig = sig[sig["practical_third"].fillna(0).astype(int) == 1]

        # --- UN Type filter (moved here to align with actor-based logic) ---
        selected_un_types = input.actors_un_type() or []
        if selected_un_types:
            sig = sig[sig["un_type"].isin(selected_un_types)]


        # Type/sub/CS combos (multi)
        combos = input.actors_type_combo() or []
        if isinstance(combos, str):
            combos = [combos]
        if combos:
            mask = pd.Series(False, index=sig.index)
            for combo in combos:
                # handle hierarchy: "Main > Sub / CS"
                combo = combo.strip()
                if ">" in combo:
                    main_type, rest = [p.strip() for p in combo.split(">", 1)]
                    subs = [p.strip() for p in rest.split("/") if p.strip()]
                else:
                    main_type = combo.strip()
                    subs = []

                cond = sig["actor_type"] == main_type
                if subs:
                    cond &= sig["sub_type"].isin(subs) | sig["cs_type"].isin(subs)
                mask |= cond
            sig = sig[mask]

        # Actor names (multi)
        names = input.actors_selected() or []
        if isinstance(names, str):
            names = [names]
        if names:
            sig = sig[sig["actor_name"].isin(names)]

        # If no actor-side filters at all, return df as-is
        if not flags and not combos and not names:
            return df

        if sig.empty:
            return pd.DataFrame(columns=df.columns)

        selected_ids = sig["AgtId"].unique()
        return df[df["AgtId"].isin(selected_ids)]

    # -------------------------------------------------------
    #  RESET + SUMMARY
    # -------------------------------------------------------
    @reactive.effect
    @reactive.event(input.reset_actors)
    def reset_actor_filters():
        ui.update_selectize("actors_type_combo", selected=[], session=session)
        ui.update_selectize("actors_selected", selected=[], session=session)
        ui.update_checkbox_group("actors_flags", selected=[], session=session)
        ui.update_checkbox("actors_third_party_only", value=False, session=session)
        ui.update_checkbox_group("actors_un_type", selected=[], session=session)


    @reactive.effect
    @reactive.event(input.reset_agreement_filters)
    def reset_agreement_filters():
        """Reset all sidebar agreement-level filters."""
        # Reset dropdowns
        ui.update_selectize("actors_region", selected=[], session=session)
        ui.update_selectize("actors_country", selected=[], session=session)
        ui.update_selectize("actors_agt_type", selected=[], session=session)
        ui.update_selectize("actors_peace_process", selected=[], session=session)
        ui.update_selectize("actors_stage", selected=[], session=session)

        # Reset year range
        yr = year_range()
        ui.update_slider("actors_year_range", value=yr, session=session)


    @render.ui
    def actors_overview():
        """
        Grey summary card:
        - "Selected actors: <list> (N)"
        - "Attributes: <flags or None>"
        - "<agreements> agreements (% of all signed agreements)"
        NOTE: % is out of unique agreements in the signatories universe.
        """
        # agreements that pass current filters (actor + general)
        df = filtered_agreements()
        sel_count = df["AgtId"].nunique() if not df.empty else 0

        # --- UN Type selections overview ---
        selected_un_types = input.actors_un_type() or []
        un_type_actors_text = None

        if selected_un_types:
            # Get all unique actor_names from signatories matching selected UN types
            matching_actors = (
                signatories[signatories["un_type"].isin(selected_un_types)]["actor_name"]
                .dropna()
                .unique()
                .tolist()
            )
            if matching_actors:
                count_text = f" ({len(matching_actors)} in total)"
                un_type_actors_text = (
                    f"Actors included under '{', '.join(selected_un_types)}' {count_text}: "
                    + ", ".join(sorted(matching_actors))
                )
            else:
                un_type_actors_text = f"No actors found for selected UN Type(s): {', '.join(selected_un_types)}"
        
        # --- Actor Type selections overview ---
        selected_combos = input.actors_type_combo() or []
        actor_type_text = None

        if selected_combos:
            matching_actors_type = []
            for combo in selected_combos:
                combo = combo.strip()
                if ">" in combo:
                    main_type, rest = [p.strip() for p in combo.split(">", 1)]
                    subs = [p.strip() for p in rest.split("/") if p.strip()]
                else:
                    main_type = combo.strip()
                    subs = []

                cond = signatories["actor_type"] == main_type
                if subs:
                    cond &= signatories["sub_type"].isin(subs) | signatories["cs_type"].isin(subs)

                matching_actors_type.extend(signatories.loc[cond, "actor_name"].dropna().unique().tolist())

            matching_actors_type = sorted(set(matching_actors_type))
            if matching_actors_type:
                count_text = f" ({len(matching_actors_type)} in total)"
                actor_type_text = (
                    f"Actors included under '{', '.join(selected_combos)}'{count_text}: "
                    + ", ".join(matching_actors_type[:8])
                    + (f" + {len(matching_actors_type)-8} more" if len(matching_actors_type) > 8 else "")
                )



        # DENOMINATOR: unique agreements that *have signatories* (signatories universe)
        #total_signed_universe = signatories["AgtId"].dropna().nunique()
        pct = (sel_count / SIGNATORY_UNIVERSE_TOTAL * 100) if SIGNATORY_UNIVERSE_TOTAL else 0

        # selected names (for display)
        names = input.actors_selected() or []
        if isinstance(names, str):
            names = [names]

        # build first line (actors or type-combos)
        display = []
        if names:
            shown = names[:3]
            if len(names) > 3:
                shown.append(f"+ {len(names) - 3} more")
            display.append(" | ".join(shown))
        else:
            combos = input.actors_type_combo() or []
            if isinstance(combos, str):
                combos = [combos]
            if combos:
                shown = combos[:3]
                if len(combos) > 3:
                    shown.append(f"+ {len(combos) - 3} more")
                display.append(" | ".join(shown))
            else:
                display.append("All")

        # count selected actors (explicit or implied)
        if names:
            n_selected_actors = len(names)
        else:
            sig = signatories.copy()

            # apply attribute flags
            flags_current = input.actors_flags() or []
            for col in ["international", "regional", "women", "practical_third"]:
                if col in sig.columns and col in flags_current:
                    sig[col] = sig[col].fillna(0).astype(int)
                    sig = sig[sig[col] == 1]

            # apply type-combo filters if any
            combos = input.actors_type_combo() or []
            if isinstance(combos, str):
                combos = [combos]
            if combos:
                mask = pd.Series(False, index=sig.index)
                for combo in combos:
                    combo = combo.strip()
                    if ">" in combo:
                        main_type, rest = [p.strip() for p in combo.split(">", 1)]
                        subs = [p.strip() for p in rest.split("/") if p.strip()]
                    else:
                        main_type = combo.strip()
                        subs = []

                    cond = sig["actor_type"] == main_type
                    if subs:
                        cond &= sig["sub_type"].isin(subs) | sig["cs_type"].isin(subs)
                    mask |= cond
                sig = sig[mask]


            n_selected_actors = sig["actor_name"].nunique()

        # attributes (checkbox flags) display
        flags = input.actors_flags() or []
        pretty = {
            "international": "International",
            "regional": "Regional",
            "women": "Women",
            "practical_third": "Third Party",
        }
        selected_flags_text = ", ".join(pretty[f] for f in flags if f in pretty) if flags else "None"

        return ui.div(
            # line 1: selected actors (N)
            ui.div(
                ui.tags.strong("Selected: ", style="color:#091f40;"),
                f"{display[0]} ",
                ui.tags.span(f"({n_selected_actors} actors selected)", style="color:#6c757d;"),
                style="margin-bottom:6px;"
            ),
            # line 2: attributes/flags
            ui.div(
                ui.tags.strong("Attributes: ", style="color:#091f40;"),
                selected_flags_text,
                style="margin-bottom:6px; font-size:0.95em; color:#495057;"
            ),
              # line 3: UN Type selections (if any)
            (ui.div(
                ui.tags.strong("UN Type selection: ", style="color:#091f40;"),
                un_type_actors_text,
                style="margin-bottom:6px; font-size:0.9em; color:#495057;"
            ) if un_type_actors_text else None),

            # line 4: Actor Type selections (if any)
            (ui.div(
                ui.tags.strong("Actor Type selection: ", style="color:#091f40;"),
                actor_type_text,
                style="margin-bottom:6px; font-size:0.9em; color:#495057;"
            ) if actor_type_text else None),

            # line 5: agreements summary
            ui.div(
                ui.tags.strong(f"{sel_count:,} agreements", style="color:#df1f36;"),
                f" ({pct:.1f}% of all signed agreements)",
                style="font-size:0.9em;color:#666;"
            ),
            style=(
                "background-color:#f8f9fa; padding:10px; border-radius:4px; "
                "border-left:3px solid #df1f36; margin-top:10px;"
            ),
        )

    # -------------------------------------------------------
    #  DATA BUILDERS FOR CHARTS
    # -------------------------------------------------------
    @reactive.calc
    def actors_map_data():
        df = filtered_agreements()
        if df.empty:
            return pd.DataFrame()

        iso_long = pd.melt(df, id_vars=["AgtId"], value_vars=["Loc1ISO", "Loc2ISO"]).dropna()
        iso_grouped = iso_long.groupby("value")["AgtId"].nunique().reset_index(name="Number of Agreements")
        geo = iso_grouped.merge(con_ent, left_on="value", right_on="iso_code", how="left").dropna(
            subset=["central_latitude", "central_longitude"]
        )
        return geo

    @reactive.calc
    def actors_time_data():
        df = filtered_agreements()
        if df.empty:
            return pd.DataFrame(columns=["year", "agreements", "total", "percentage"])

        yearly = df.groupby("year")["AgtId"].nunique().reset_index(name="agreements")
        merged = signatory_year_totals.merge(yearly, on="year", how="left").fillna({"agreements": 0})
        merged["percentage"] = merged["agreements"] / merged["total"] * 100
        return merged[["year", "agreements", "total", "percentage"]]

    @reactive.calc
    def actors_stage_data():
        """
        Stage distribution for both the selected set and the whole signatory universe.
        - percentage_selected: % of selected agreements at each stage (denom = selected unique AgtId)
        - percentage_all: % of all signatory agreements at each stage (denom = SIGNATORY_UNIVERSE_TOTAL)
        Also keeps selected counts for Count mode.
        """
        df_sel = filtered_agreements()
        all_stages = pd.DataFrame({"stage_label": stage_order})

        # --- Selected set (counts + % within selection) ---
        if df_sel.empty:
            sel_counts = pd.DataFrame({"stage_label": [], "count_selected": []})
            denom_sel = 1
        else:
            sel_counts = (
                df_sel.groupby("stage_label")["AgtId"]
                .nunique()
                .reset_index(name="count_selected")
            )
            denom_sel = df_sel["AgtId"].nunique() or 1

        # --- Global universe (grey) ---
        # Use signatories only (not all pax); ensure unique AgtId x stage_label first.
        all_counts = (
            signatories[["AgtId", "stage_label"]]
            .dropna(subset=["stage_label"])
            .drop_duplicates()
            .groupby("stage_label")["AgtId"]
            .nunique()
            .reset_index(name="count_all")
        )
        denom_all = SIGNATORY_UNIVERSE_TOTAL or 1

        # --- Merge and compute percentages ---
        out = (
            all_stages
            .merge(sel_counts, on="stage_label", how="left")
            .merge(all_counts, on="stage_label", how="left")
            .fillna({"count_selected": 0, "count_all": 0})
        )
        out["count_selected"] = out["count_selected"].astype(int)
        out["count_all"] = out["count_all"].astype(int)

        out["percentage_selected"] = out["count_selected"] / denom_sel * 100
        out["percentage_all"] = out["count_all"] / denom_all * 100

        return out

    @reactive.calc
    def actors_cosign_split_data():
        """Co-sign partners (Top N) split Party vs Third-party, excluding self if explicitly chosen."""
        df = filtered_agreements()
        req_n = input.actors_top_n() or 15
        if df.empty or signatories.empty:
            return {"party": pd.DataFrame(), "third": pd.DataFrame()}

        names = input.actors_selected() or []
        if isinstance(names, str):
            names = [names]

        agt_ids = df["AgtId"].unique()
        sigs = signatories[signatories["AgtId"].isin(agt_ids)].copy()
        if len(names) == 1:
            sigs = sigs[sigs["actor_name"] != names[0]]  # exclude self only in single-actor context

        sigs["signatory_type"] = np.where(sigs["practical_third"] == 1, "third", "party")

        def topn(d):
            if d.empty:
                return d
            avail = d["actor_name"].nunique()
            n = min(req_n, max(1, avail))  # allow fewer than 5 gracefully
            return (
                d.groupby("actor_name")["AgtId"].nunique().reset_index(name="count")
                .sort_values("count", ascending=True).tail(int(n))
            )

        return {"party": topn(sigs[sigs["signatory_type"] == "party"]),
                "third": topn(sigs[sigs["signatory_type"] == "third"])}

    @reactive.calc
    def actors_topics_data():
        """Aggregate topics by selected level, robust to column naming in pax_topics."""
        df = filtered_agreements()
        if df.empty or pax_topics.empty:
            return pd.DataFrame(columns=["level", "count"])

        # Resolve column names robustly
        tcols = {c.lower(): c for c in pax_topics.columns}
        agt_col = tcols.get("agtid") or tcols.get("agt_id") or ("AgtId" if "AgtId" in pax_topics.columns else None)
        val_col = tcols.get("value")
        cat_col = tcols.get("category") or ("category" if "category" in pax_topics.columns else None)
        iss_col = tcols.get("issue_label") or tcols.get("issue") or ("issue_label" if "issue_label" in pax_topics.columns else None)

        if agt_col is None:
            return pd.DataFrame(columns=["level", "count"])

        tt = pax_topics[pax_topics[agt_col].isin(df["AgtId"].unique())].copy()
        if val_col and val_col in tt.columns:
            tt = tt[tt[val_col] > 0]

        level = input.actors_topics_level()
        group_col = cat_col if level == "Category" else iss_col
        if not group_col or group_col not in tt.columns:
            return pd.DataFrame(columns=["level", "count"])

        out = (
            tt.groupby(group_col)[agt_col]
            .nunique()
            .reset_index(name="count")
            .rename(columns={group_col: "level"})
            .sort_values("count", ascending=True)
        )
        return out

    @reactive.calc
    def actors_table_data():
        """Tabular list of agreements for selected actors (with key fields)."""
        df = filtered_agreements()
        if df.empty:
            return pd.DataFrame(columns=["AgtId", "Agreement", "Year", "Country", "Stage", "Process"])

        base = df[["AgtId", "agt_name", "year", "stage_label", "PPName", "Con", "PAX_Hyperlink"]].drop_duplicates()
        base.rename(
            columns={
                "agt_name": "Agreement",
                "year": "Year",
                "stage_label": "Stage",
                "PPName": "Process",
                "Con": "Country",
                "PAX_Hyperlink":"Link"
            },
            inplace=True,
        )
         # Convert Link to clickable HTML
        base["Link"] = base["Link"].apply(
            lambda url: f'<a href="{url}" target="_blank">View on PA-X</a>' if isinstance(url, str) and url.startswith("http") else ""
        )
        return base.sort_values(["Year", "Agreement"], ascending=[False, True])

    # -------------------------------------------------------
    #  FIGURES
    # -------------------------------------------------------
    def make_actors_map_figure():
        geo = actors_map_data()
        if geo.empty:
            fig = go.Figure()
            fig.add_annotation(text="No geographic data available for selected actors",
                               xref="paper", yref="paper", x=0.5, y=0.5,
                               showarrow=False, font=dict(size=16))
            fig.update_layout(height=500)
            return fig

        fig = px.scatter_geo(
            geo,
            lat="central_latitude",
            lon="central_longitude",
            hover_name="name",
            size="Number of Agreements",
            hover_data={"Number of Agreements": True, "central_latitude": False, "central_longitude": False},
            projection="natural earth",
        )
        fig.update_layout(
            title={"text": f"Geographic Distribution of Agreements signed by {actor_title_prefix()}", "x": 0.5, "xanchor": "center", "font": {"size": 18}},
            height=500,
        )
        fig.update_traces(
            marker=dict(color="#091F40", line=dict(width=1, color="white"), sizemin=5, sizemode="area")
        )
        fig.update_geos(landcolor="#f0e8d9", showframe=False, showcoastlines=True, coastlinecolor="white")
        return fig

    def make_actors_over_time_figure():
        time = actors_time_data()
        fig, ax = plt.subplots(figsize=(14, 8))
        if time.empty:
            ax.text(0.5, 0.5, "No data available", ha="center", va="center", transform=ax.transAxes)
            return fig

        if input.actors_time_mode() == "Percentage":
            y = time["percentage"]; color = "#7b8ad6"; ylab = "Percentage of Signatory Agreements"
            labels = time["percentage"].round(1).astype(str) + "%"
        else:
            y = time["agreements"]; color = "#091f40"; ylab = "Number of Agreements"
            labels = time["agreements"].astype(int).astype(str)

        ax.plot(time["year"], y, marker="o", color=color, linewidth=2, markersize=6)
        for x, yy, label in zip(time["year"], y, labels):
            if yy > 0:
                ax.text(x, yy + max(y) * 0.02, label, ha="center", va="bottom", fontsize=8, fontweight="bold")

        ax.set_xlabel("Year", fontsize=12)
        ax.set_ylabel(ylab, fontsize=12)
        ax.set_title(f"Number of Agreements signed by {actor_title_prefix()} per Year", fontsize=16, fontweight="bold", pad=20) #yhere
        ax.grid(alpha=0.3)
        plt.ylim(0, max(y) * 1.15 if len(y) else 1)
        plt.tight_layout()
        return fig

    def make_actors_stage_figure():
        """
        Count mode: navy bars show selected-set counts per stage.
        Percentage mode: side-by-side bars:
        - grey = % of all signatory agreements at that stage (global universe)
        - navy = % of selected-set agreements at that stage (within-selection %)
        """
        data = actors_stage_data()
        fig, ax = plt.subplots(figsize=(14, 8))

        if input.actors_stage_mode() == "Percentage":
            x = np.arange(len(data))
            width = 0.38

            bars_all = ax.bar(
                x - width / 2,
                data["percentage_all"],
                width,
                label=f"% of All Agreements at Stage",
                color="#cccccc",
            )
            bars_sel = ax.bar(
                x + width / 2,
                data["percentage_selected"],
                width,
                label=f"% of Agreements signed by {actor_title_prefix()} at Stage",
                color="#091f40",
            )

            # Labels
            for b, val in zip(bars_all, data["percentage_all"]):
                if val > 0:
                    ax.text(b.get_x() + b.get_width()/2, val + 0.8, f"{val:.0f}%", ha="center", va="bottom", fontsize=9, fontweight="bold")
            for b, val in zip(bars_sel, data["percentage_selected"]):
                if val > 0:
                    ax.text(b.get_x() + b.get_width()/2, val + 0.8, f"{val:.0f}%", ha="center", va="bottom", fontsize=9, fontweight="bold")

            ax.set_xticks(x)
            ax.set_xticklabels(data["stage_label"], rotation=45, ha="right")
            ax.set_ylabel("Percentage of Agreements")
            ymax = max(data["percentage_all"].max(), data["percentage_selected"].max())
            ax.set_ylim(0, ymax * 1.25 if ymax > 0 else 1)
            ax.legend(bbox_to_anchor=(0.5, 1), loc="lower center", ncol=2, frameon=False, fontsize=10)

        else:
            # Count mode (selected only)
            bars = ax.bar(data["stage_label"], data["count_selected"], color="#091f40")
            for b, val in zip(bars, data["count_selected"]):
                if val > 0:
                    ax.text(b.get_x() + b.get_width()/2, val + 0.3, f"{int(val)}", ha="center", va="bottom", fontsize=9, fontweight="bold")
            plt.setp(ax.get_xticklabels(), rotation=45, ha="right")
            ax.set_ylabel("Number of Agreements")
            ymax = data["count_selected"].max()
            ax.set_ylim(0, ymax * 1.25 if ymax > 0 else 1)

        ax.set_xlabel("Stage", fontsize=12)
        ax.set_title(
            f"Agreements signed by {actor_title_prefix()}, by Stage of Process",
            fontsize=16,
            fontweight="bold",
            pad=20,
            y=1.05,
        )
        plt.tight_layout()
        return fig


    @reactive.calc
    def actors_peace_process_data():
        """
        Top N peace processes for selected actors.
        If actor signs in only a few processes, adaptively shrink N
        so empty bars are not shown.
        """
        df = filtered_agreements()
        mode = input.actors_pp_view_mode()  # "count", "percentage", or "stage"
        n = input.actors_top_processes() or 20
        if df.empty or "PPName" not in df.columns:
            return pd.DataFrame()

        sel = (
            df[["AgtId", "PPName", "stage_label"]]
            .dropna(subset=["PPName"])
            .assign(PPName=lambda d: d["PPName"].astype(str).str.strip())
            .drop_duplicates()
        )

        universe = (
            signatories[["AgtId", "PPName"]]
            .dropna(subset=["PPName"])
            .assign(PPName=lambda d: d["PPName"].astype(str).str.strip())
            .drop_duplicates()
            .groupby("PPName")["AgtId"]
            .nunique()
            .reset_index(name="total")
        )

        selected_pp = (
            sel.groupby("PPName")["AgtId"]
            .nunique()
            .reset_index(name="count")
        )

        merged = universe.merge(selected_pp, on="PPName", how="left").fillna({"count": 0})
        merged["percentage"] = merged["count"] / merged["total"] * 100

        # --- NEW: filter to processes with at least one agreement ---
        merged = merged[merged["count"] > 0]

        # --- NEW: adaptive N if fewer valid processes ---
        available_n = len(merged)
        n = min(n, available_n if available_n > 0 else 5)

        # --- Stage breakdown ---
        if mode == "stage":
            stage_data = (
                sel.groupby(["PPName", "stage_label"])["AgtId"]
                .nunique()
                .reset_index(name="count")
            )
            stage_pivot = (
                stage_data.pivot(index="PPName", columns="stage_label", values="count")
                .fillna(0)
            )
            # keep top N by total across stages
            totals = stage_pivot.sum(axis=1).sort_values(ascending=True).tail(int(n))
            return stage_pivot.loc[totals.index]

        # Otherwise percentage or count single-series
        elif mode == "percentage":
            out = merged.sort_values("percentage", ascending=True).tail(int(n))
            return out[["PPName", "percentage"]]
        else: #count
            out = merged.sort_values("count", ascending=True).tail(int(n))
            return out[["PPName", "count"]]


    def make_actors_peace_process_figure():
        df = actors_peace_process_data()
        mode = input.actors_pp_view_mode()
        fig, ax = plt.subplots(figsize=(14, 10))

        if df.empty:
            ax.text(0.5, 0.5, "No data available", ha="center", va="center", transform=ax.transAxes)
            return fig

        # Determine how many processes are actually displayed
        displayed_n = df.shape[0]
        slider_n = input.actors_top_processes() or 20
        show_n_text = f"(All {displayed_n} Processes)" if displayed_n < slider_n else f"(Top {slider_n})"

        # === Without stage breakdown ===
        if mode != "stage":
            bars = ax.barh(df["PPName"], df.iloc[:, -1], color="#091f40")

            for bar in bars:
                width = bar.get_width()
                label = f"{width:.0f}%" if mode == "percentage" else f"{int(width)}"
                ax.text(width + max(df.iloc[:, -1]) * 0.01,
                        bar.get_y() + bar.get_height() / 2,
                        label,
                        ha="left", va="center", fontsize=9)

            xlabel = "Percentage of Agreements" if mode == "percentage" else "Number of Agreements"
            ax.set_xlabel(xlabel)
            ax.set_ylabel("Peace Process")
            ax.set_title(
                f"Peace Processes with {actor_title_prefix()} {show_n_text}",
                fontsize=16, fontweight="bold"
            )

        # === With stage breakdown ===
        else:
            stage_pivot = df
            colors = [stage_color_map.get(s, "#cccccc") for s in stage_pivot.columns]
            stage_pivot.plot(kind="barh", stacked=True, color=colors, ax=ax, width=0.8)

            for container in ax.containers:
                # Label only bars with a nonzero width
                labels = [f"{int(w)}" if w > 0 else "" for w in container.datavalues]
                ax.bar_label(container, labels=labels, label_type="center", color="white", fontsize=8)

            totals = stage_pivot.sum(axis=1)
            for i, total in enumerate(totals):
                if total > 0:
                    ax.text(total + totals.max() * 0.01, i, f"{int(total)}", va="center", fontsize=9, fontweight="bold")

            ncol = min(4, len(stage_pivot.columns))
            ax.legend(title="Stage", bbox_to_anchor=(0.5, 1), loc="lower center", ncol=ncol, frameon=False)
            ax.set_xlabel("Number of Agreements")
            ax.set_ylabel("Peace Process")
            ax.set_title(f"Peace Processes by Stage of Process {show_n_text}",
                        fontsize=16, fontweight="bold", pad=15, y=1.1)

        plt.tight_layout()
        return fig



    def make_party_cosign_figure():
        data = actors_cosign_split_data()["party"]
        fig, ax = plt.subplots(figsize=(12, 8))
        if data.empty:
            ax.text(0.5, 0.5, "No co-sign party data", ha="center", va="center", transform=ax.transAxes)
            return fig
        bars = ax.barh(data["actor_name"], data["count"], color="#091f40")
        for b in bars:
            w = b.get_width()
            ax.text(w + max(data["count"]) * 0.01, b.get_y() + b.get_height()/2, f"{int(w)}",
                    ha="left", va="center", fontsize=9, fontweight="bold")
        ax.set_xlim(0, data["count"].max() * 1.15)
        ax.set_xlabel("Number of Co-signed Agreements")
        n = input.actors_top_n() or 15
        ax.set_title(f"Top {int(n)} Party Co-signatories with {actor_title_prefix()}", fontsize=16, fontweight="bold")
        plt.tight_layout()
        return fig

    def make_third_cosign_figure():
        data = actors_cosign_split_data()["third"]
        fig, ax = plt.subplots(figsize=(12, 8))
        if data.empty:
            ax.text(0.5, 0.5, "No co-sign third-party data", ha="center", va="center", transform=ax.transAxes)
            return fig
        bars = ax.barh(data["actor_name"], data["count"], color="#df1f36")
        for b in bars:
            w = b.get_width()
            ax.text(w + max(data["count"]) * 0.01, b.get_y() + b.get_height()/2, f"{int(w)}",
                    ha="left", va="center", fontsize=9, fontweight="bold")
        ax.set_xlim(0, data["count"].max() * 1.15)
        ax.set_xlabel("Number of Co-signed Agreements")
        n = input.actors_top_n() or 15
        ax.set_title(f"Top {int(n)} Third-Party Co-signatories with {actor_title_prefix()}", fontsize=16, fontweight="bold")
        plt.tight_layout()
        return fig

    def make_actor_topics_figure():
        data = actors_topics_data()
        if data.empty:
            fig, ax = plt.subplots(figsize=(10, 5))
            ax.text(0.5, 0.5, "No topic data for selected actors", ha="center", va="center", transform=ax.transAxes)
            return fig

        # Adjust figure height dynamically
        n_topics = len(data)
        base_height = 0.4
        fig_height = max(6, n_topics * base_height)
        fig, ax = plt.subplots(figsize=(14, fig_height))

        bars = ax.barh(data["level"], data["count"], color="#091f40")
        for b in bars:
            w = b.get_width()
            ax.text(w + max(data["count"]) * 0.01,
                    b.get_y() + b.get_height()/2,
                    f"{int(w)}",
                    ha="left", va="center", fontsize=9, fontweight="bold")

        ax.set_xlim(0, data["count"].max() * 1.15)
        ax.set_title(f"Topics in Agreements signed by {actor_title_prefix()}", fontsize=16, fontweight="bold")
            # --- Fix spacing above and below bars ---
        ax.margins(y=0)  # remove default top/bottom padding
        ax.set_ylim(-0.5, len(data) - 0.5)  # exact fit to bars
        ax.set_xlabel("Number of agreements")
        plt.subplots_adjust(top=0.96, bottom=0.06, left=0.32, right=0.95)
        plt.tight_layout(rect=[0, 0, 1, 1])
        return fig


    def make_actor_topics_radial_figure():
        data = actors_topics_data()
        if data.empty:
            fig, ax = plt.subplots(figsize=(8, 6))
            ax.text(0.5, 0.5, "No topic data available", ha="center", va="center", transform=ax.transAxes)
            return fig

        fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
        categories = data["level"]
        values = data["count"]
        angles = np.linspace(0, 2 * np.pi, len(categories), endpoint=False).tolist()

        values = np.concatenate((values, [values[0]]))
        angles += angles[:1]

        ax.plot(angles, values, color="#091f40", linewidth=2)
        ax.fill(angles, values, color="#091f40", alpha=0.4)
        ax.set_yticklabels([])
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(categories, fontsize=9)
        ax.set_title(f"Topics in Agreements signed by {actor_title_prefix()}", fontsize=14, fontweight="bold", pad=20)
        plt.tight_layout()
        return fig

    def make_actors_agreement_type_pie():
        df = filtered_agreements()
        if df.empty:
            fig, ax = plt.subplots(figsize=(8, 8))
            ax.text(0.5, 0.5, "No data available", ha="center", va="center", transform=ax.transAxes)
            return fig

        type_counts = df["agt_type"].value_counts().reset_index()
        type_counts.columns = ["Type", "Count"]

        colors = [type_color_map.get(t, "#cccccc") for t in type_counts["Type"]]
        fig, ax = plt.subplots(figsize=(8, 6))
        ax.pie(
            type_counts["Count"],
            labels=type_counts["Type"],
            autopct="%1.0f%%",
            colors=colors,
            startangle=140,
            textprops={"fontsize": 9},
        )
        ax.set_title(f"Types of Agreements signed by {actor_title_prefix()}", fontsize=14, fontweight="bold")
        plt.tight_layout()
        return fig

    # -------------------------------------------------------
    #  RENDER: MAP (Plotly widget) + PLOTS
    # -------------------------------------------------------
    @render_widget
    def actors_map():
        return make_actors_map_figure()

    @render.plot
    def actors_over_time():
        return make_actors_over_time_figure()

    @render.plot
    def actors_by_stage():
        return make_actors_stage_figure()

    @render.plot
    def actors_by_peace_process():
        return make_actors_peace_process_figure()

    @render.plot
    def actors_party_cosign():
        return make_party_cosign_figure()

    @render.plot
    def actors_third_cosign():
        return make_third_cosign_figure()
    
    # === Actual Topics chart (Matplotlib) ===
    @render.plot
    def actors_topics_chart():
        """Render the main topics chart (Matplotlib)."""
        return make_actor_topics_figure()

    # @render.plot
    # def actors_topics_chart():
    #     return make_actor_topics_figure()
    # === Dynamic container for Topics chart ===
    @render.ui
    def actors_topics_plot_container():
        """Dynamically adjust the visible height of the topics plot container."""
        data = actors_topics_data()
        # Estimate display height based on number of topics (each ~35px)
        if data is None or data.empty:
            dynamic_height = 700
        else:
            dynamic_height = max(750, len(data) * 35)
        # Return a UI container that sets the height dynamically
        return ui.output_plot("actors_topics_chart", height=f"{dynamic_height}px")

   
   
    @render.plot
    def actors_topics_radial_chart():
        return make_actor_topics_radial_figure()
    @render.ui
    def actors_topics_radial_plot_container():
        """Dynamically adjust the visible height of the radial topics chart container."""
        data = actors_topics_data()
        if data is None or data.empty:
            dynamic_height = 700
        else:
            # Radial charts need a bit more vertical space to look balanced
            dynamic_height = max(800, len(data) * 6)
        return ui.output_plot("actors_topics_radial_chart", height=f"{dynamic_height}px")

    @render.plot
    def actors_agreement_type_pie():
        return make_actors_agreement_type_pie()

    # -------------------------------------------------------
    #  DOWNLOADS (PNG + CSV)
    # -------------------------------------------------------
    @render.download(filename="actors_over_time.png")
    def actors_over_time_png():
        return export_with_branding(
            make_actors_over_time_figure,
            filter_text_fn=get_actor_filter_text,
            data_version_fn=get_data_version,
            load_data_fn=lambda: {"pax": pax},
            logo_position=ACTORS_LOGO_POSITION,
            filter_text_position=(0.5, 0.99),
            version_position=ACTORS_VERSION_POSITION,
        )

    @render.download(filename="actors_over_time.csv")
    def actors_over_time_csv():
        df = actors_time_data()
        return io.BytesIO(df.to_csv(index=False).encode())

    @render.download(filename="actors_by_stage.png")
    def actors_by_stage_png():
        return export_with_branding(
            make_actors_stage_figure,
            filter_text_fn=get_actor_filter_text,
            data_version_fn=get_data_version,
            load_data_fn=lambda: {"pax": pax},
            logo_position=ACTORS_LOGO_POSITION,
            filter_text_position=(0.5, 0.99),
            version_position=ACTORS_VERSION_POSITION,
        )

    @render.download(filename="actors_by_stage.csv")
    def actors_by_stage_csv():
        df = actors_stage_data()
        return io.BytesIO(df.to_csv(index=False).encode())

    @render.download(filename="actors_map.csv")
    def actors_map_csv():
        geo = actors_map_data()
        if not geo.empty:
            export_data = geo[["value", "name", "Number of Agreements", "central_latitude", "central_longitude"]].copy()
            export_data.columns = ["ISO_Code", "Country_Name", "Number_of_Agreements", "Latitude", "Longitude"]
        else:
            export_data = pd.DataFrame(columns=["ISO_Code", "Country_Name", "Number_of_Agreements", "Latitude", "Longitude"])
        return io.BytesIO(export_data.to_csv(index=False).encode("utf-8"))

    @render.download(filename="actors_peace_process.png")
    def actors_peace_process_png():
        return export_with_branding(
            make_actors_peace_process_figure,
            filter_text_fn=get_actor_filter_text,
            data_version_fn=get_data_version,
            load_data_fn=lambda: {"pax": pax},
            logo_position=(1, 1, 0.075, 0.075),
            filter_text_position=(0.2, 1),
            version_position=ACTORS_VERSION_POSITION,
        )

    @render.download(filename="actors_peace_process.csv")
    def actors_peace_process_csv():
        df = actors_peace_process_data()
        return io.BytesIO(df.to_csv(index=False).encode())

    @render.download(filename="actors_party_cosign.png")
    def actors_party_cosign_png():
        return export_with_branding(
            make_party_cosign_figure,
            filter_text_fn=get_actor_filter_text,
            data_version_fn=get_data_version,
            load_data_fn=lambda: {"pax": pax},
            logo_position=ACTORS_LOGO_POSITION,
            filter_text_position=(0.1, 0.0035),
            version_position=ACTORS_VERSION_POSITION,
        )

    @render.download(filename="actors_third_cosign.png")
    def actors_third_cosign_png():
        return export_with_branding(
            make_third_cosign_figure,
            filter_text_fn=get_actor_filter_text,
            data_version_fn=get_data_version,
            load_data_fn=lambda: {"pax": pax},
            logo_position=ACTORS_LOGO_POSITION,
            filter_text_position=(0.1, 0.0035),
            version_position=ACTORS_VERSION_POSITION,
        )

    @render.download(filename="actors_party_cosign.csv")
    def actors_party_cosign_csv():
        df = actors_cosign_split_data()["party"]
        return io.BytesIO(df.to_csv(index=False).encode())

    @render.download(filename="actors_third_cosign.csv")
    def actors_third_cosign_csv():
        df = actors_cosign_split_data()["third"]
        return io.BytesIO(df.to_csv(index=False).encode())

    @render.download(filename="actors_topics.png")
    def actors_topics_png():
        return export_with_branding(
            make_actor_topics_figure,
            filter_text_fn=get_actor_filter_text,
            data_version_fn=get_data_version,
            load_data_fn=lambda: {"pax": pax},
            logo_position=ACTORS_LOGO_POSITION,
            filter_text_position=(0.5, 1.1),
            version_position=ACTORS_VERSION_POSITION,
        )

    @render.download(filename="actors_topics.csv")
    def actors_topics_csv():
        df = actors_topics_data()
        return io.BytesIO(df.to_csv(index=False).encode())
    
    @render.download(filename="actors_radial_topics.png") 
    def actors_topics_radial_png():
        return export_with_branding(
            make_actor_topics_radial_figure,
            filter_text_fn=get_actor_filter_text,
            data_version_fn=get_data_version,
            load_data_fn=lambda: {"pax": pax},
            logo_position=(0.95, 0.9, 0.075, 0.075),
            filter_text_position=ACTORS_FILTER_TEXT_POSITION,
            version_position=ACTORS_VERSION_POSITION,
        )

    @render.download(filename="actors_radial_topics.csv")
    def actors_topics_radial_csv():
        df = actors_topics_data()
        return io.BytesIO(df.to_csv(index=False).encode())

    @render.download(filename="actors_agreement_type_pie.png")
    def actors_agreement_type_pie_png():
        return export_with_branding(
            make_actors_agreement_type_pie,
            filter_text_fn=get_actor_filter_text,
            data_version_fn=get_data_version,
            load_data_fn=lambda: {"pax": pax},
            logo_position=(0.95, 0.92, 0.075, 0.075),
            filter_text_position=ACTORS_FILTER_TEXT_POSITION,
            version_position=ACTORS_VERSION_POSITION,
        )

    @render.download(filename="actors_agreement_type_pie.csv")
    def actors_agreement_type_pie_csv():
        df = filtered_agreements()[["agt_type", "AgtId"]]
        df = df.groupby("agt_type")["AgtId"].nunique().reset_index(name="count")
        return io.BytesIO(df.to_csv(index=False).encode())

    # -------------------------------------------------------
    #  TABLE (styled)
    # -------------------------------------------------------
    @render.ui
    def actors_table():
        df = actors_table_data()
        if df.empty:
            return ui.div("No agreements found for the selected filters.",
                        style="color: #6c757d; font-style: italic;")

        # Build a clean HTML table manually
        table_html = df.to_html(
            index=False,
            classes="table table-striped table-hover align-middle",
            justify="center",
            border=0,
            escape=False  # <-- allows clickable links in the next step
        )

        # Wrap with Bootstrap style
        return ui.HTML(f"""
            <div style='overflow-x: auto;'>
                <style>
                    table {{
                        width: 100%;
                        border-collapse: collapse;
                        font-size: 0.9rem;
                    }}
                    th {{
                        background-color: #f8f9fa;
                        color: #091f40;
                        text-align: center;
                        padding: 8px;
                        border-bottom: 2px solid #dee2e6;
                    }}
                    td {{
                        text-align: center;
                        padding: 8px;
                        vertical-align: middle;
                    }}
                    tr:nth-child(even) {{ background-color: #f9f9f9; }}
                </style>
                {table_html}
            </div>
        """)
    
    ## --- TABLE EXPORTS -----

    @output
    @render.download(filename="actors_table.csv")
    def actors_table_csv():
        df = actors_table_data().copy()
        # Strip out HTML tags before export
        df["Link"] = df["Link"].str.replace(r"<.*?>", "", regex=True)
        csv_bytes = df.to_csv(index=False).encode("utf-8")
        return BytesIO(csv_bytes)


    # @output
    # @render.download(filename="actors_table.docx")
    # def actors_table_docx():
    #     df = actors_table_data().copy()
    #     df["Link"] = df["Link"].str.replace(r"<.*?>", "", regex=True)

    #     doc = Document()
    #     doc.add_heading("Agreements Table", level=1)

    #     # Create table
    #     table = doc.add_table(rows=1, cols=len(df.columns))
    #     hdr_cells = table.rows[0].cells
    #     for i, col in enumerate(df.columns):
    #         hdr_cells[i].text = col

    #     for _, row in df.iterrows():
    #         cells = table.add_row().cells
    #         for i, val in enumerate(row):
    #             cells[i].text = str(val)

    #     buf = BytesIO()
    #     doc.save(buf)
    #     buf.seek(0)
    #     return buf
    @output
    @render.download(filename="actors_table.docx")
    def actors_table_docx():
        """Export the filtered agreements table to a styled Word document with working hyperlinks."""
        df = actors_table_data()
        doc = Document()

        # === DYNAMIC TITLE ===
        selected_actors = input.actors_selected() or []
        if isinstance(selected_actors, str):
            selected_actors = [selected_actors]
        if selected_actors:
            actor_list = ", ".join(selected_actors)
            title_text = f"Agreements signed by {actor_list}"
        else:
            title_text = "Agreements signed by selected actors"

        title_para = doc.add_paragraph(title_text)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.runs[0]
        run.font.name = "Montserrat"
        run.font.size = Pt(16)
        run.bold = True

        doc.add_paragraph("")  # spacer

        # === TABLE ===
        if df.empty:
            doc.add_paragraph("No agreements available for the selected filters.")
        else:
            df = df.rename(
                columns={
                    "Agreement": "Agreement Title",
                    "Year": "Year",
                    "Country": "Country",
                    "Stage": "Stage of Process",
                    "Process": "Peace Process",
                    "Link": "PA-X Link",
                }
            )

            # Create table
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Table Grid"

            # --- Header row ---
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df.columns):
                p = hdr_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(col_name)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.bold = True

                # Light gray header shading
                tcPr = hdr_cells[i]._element.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "D9D9D9")
                tcPr.append(shd)

            # --- Data rows ---
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, col_name in enumerate(df.columns):
                    val = str(row[col_name]) if not pd.isna(row[col_name]) else ""
                    p = row_cells[i].paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)

                    # Ensure consistent text styling
                    for r in p.runs:
                        r.bold = False
                        r.font.name = "Calibri"
                        r.font.size = Pt(10)

                    # --- Handle hyperlink column ---
                    if col_name == "PA-X Link" and val.startswith("<a "):
                        match = re.search(r'href="([^"]+)"[^>]*>(.*?)<', val)
                        if match:
                            href, label = match.groups()
                            add_hyperlink(p, href, label, font_size=10)
                        else:
                            run = p.add_run(val)
                            run.font.name = "Calibri"
                            run.font.size = Pt(10)
                            run.bold = False
                    else:
                        run = p.add_run(val)
                        run.font.name = "Calibri"
                        run.font.size = Pt(10)
                        run.bold = False

            # Fit table neatly
            table.autofit = True
            for row in table.rows:
                for cell in row.cells:
                    cell.width = Inches(1.2)

        # === FOOTER INFO ===
        doc.add_paragraph("")  # spacer before footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = ""

        # Pull metadata from your existing helper functions
        try:
            filters_text = get_actor_filter_text() or "No filters applied"
        except Exception:
            filters_text = "No filters applied"

        try:
            version_text = get_data_version() or "Unknown"
        except Exception:
            version_text = "Unknown"

        footer_text = f"Filters applied: {filters_text}\nPA-X Database v{version_text}"

        footer_run = footer_para.add_run(footer_text)
        footer_run.font.name = "Calibri"
        footer_run.font.size = Pt(9)
        footer_run.italic = True
        footer_run.font.color.rgb = None

        # Save buffer for download
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


    # === Utility: make Word hyperlinks clickable ===
    def add_hyperlink(paragraph, url, text, color="0000EE", underline=True, font_size=10):
        """Add a clickable hyperlink to a Word paragraph with consistent font size."""
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Blue + underline
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
        if underline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)

        # Font size (half-points)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size * 2)))
        rPr.append(sz)

        new_run.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        new_run.append(t)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        return hyperlink